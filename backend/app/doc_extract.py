"""Text out of a document the operator handed us (roadmap #22a).

Why this exists: the composer refused every PDF. `addFiles` read each
non-image with `File.text()` and rejected anything whose first bytes decoded
with replacement characters — which is precisely what a PDF does. That
refusal was honest and correct; there was simply nothing that could read the
file. So "here, look at this document" was answerable for .txt and .md and
for nothing anyone actually gets sent.

Two rules shape this module, and both are the mechanical-over-prompts rule
applied to a parser:

1. **Extraction failure is stated, never silently thin.** A scanned PDF is
   images of text: every library returns pages of empty strings for it, and
   an empty string inlined into a turn reads to the model as "the document
   said nothing." It would then answer about a document it never saw. So a
   result that is empty or nearly so comes back as a REFUSAL carrying the
   reason, and the caller shows that instead of pretending.

2. **Truncation announces itself.** Same argument as `_cap_result` in the
   runner: a model cannot flag a gap it was never shown.

Nothing here trusts the client's mime — a name and a mime are both operator
input. The sniff is on the file's own leading bytes, and the mime is only a
tiebreak.
"""

import io
import logging

log = logging.getLogger(__name__)

# Per FILE, after extraction. The 20 MB whole-message cap in router_chat
# bounds the upload; this bounds what a single document may spend of the
# turn's context. A 300-page PDF is perfectly legal to attach and would
# otherwise evict the entire conversation.
MAX_CHARS = 60_000

# Below this, an "extraction" is indistinguishable from a failed one. A
# 40-page scanned contract and a genuinely one-line file both come back
# tiny; only the first is a lie, but neither is worth guessing about, so
# both get told the truth about what came out.
_MIN_USEFUL_CHARS = 16


class Unsupported(Exception):
    """No extractor for this file. Carries operator-facing wording."""


class Unextractable(Exception):
    """There is an extractor, and it got nothing usable out."""


def _sniff(data: bytes, name: str, mime: str) -> str:
    """What this file actually IS — bytes first, name and mime as tiebreaks."""
    if data[:5] == b"%PDF-":
        return "pdf"
    # docx/odt/etc are zip containers; the member list distinguishes them
    if data[:2] == b"PK":
        lowered = name.lower()
        if lowered.endswith(".docx") or "wordprocessingml" in mime:
            return "docx"
        return "zip"
    return "text"


def _pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover - deployment error, not input
        raise Unsupported("PDF support is not installed on this server") from e
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:
        raise Unextractable(f"the PDF could not be opened ({e})") from e
    if getattr(reader, "is_encrypted", False):
        # an empty decrypt succeeds on many "protected" PDFs; try it once
        try:
            reader.decrypt("")
        except Exception:
            raise Unextractable("the PDF is password-protected") from None
    parts = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            log.debug("page %d of a PDF would not extract", i, exc_info=True)
            parts.append("")
    text = "\n\n".join(p.strip() for p in parts if p.strip())
    if len(text) < _MIN_USEFUL_CHARS:
        # No claim about OCR here any more. This used to end "which needs OCR
        # this server does not have" — true when written, false the moment
        # tesseract landed in the image, and it is the kind of sentence
        # nobody revisits. Whether OCR is available is asked at the point of
        # use (ocr.available(), which looks for the binaries) and appended by
        # extract_best, so the two can never disagree.
        raise Unextractable(
            f"the PDF has {len(reader.pages)} page(s) but no extractable text "
            f"layer — it is most likely a scan")
    return text


def _docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as e:  # pragma: no cover
        raise Unsupported("Word support is not installed on this server") from e
    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as e:
        raise Unextractable(f"the Word file could not be opened ({e})") from e
    # tables carry the content in plenty of real documents, and paragraphs
    # alone silently drop them
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    text = "\n".join(c for c in (s.strip() for s in chunks) if c)
    if len(text) < _MIN_USEFUL_CHARS:
        raise Unextractable("the Word file contains no readable text")
    return text


def _plain(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    # the same replacement-character test the composer used, kept because it
    # is the thing that correctly identifies "this is not text at all"
    if text[:2000].count("�") > 4:
        raise Unsupported("this looks like a binary file, not a document")
    if len(text.strip()) < _MIN_USEFUL_CHARS and not text.strip():
        raise Unextractable("the file is empty")
    return text


def _cap(text: str, name: str) -> str:
    if len(text) <= MAX_CHARS:
        return text
    log.info("attachment %r truncated: %d of %d chars", name,
             MAX_CHARS, len(text))
    return (text[:MAX_CHARS] + f"\n\n[TRUNCATED — you were shown "
            f"{MAX_CHARS:,} of {len(text):,} characters of "
            f"{name or 'this file'}. The rest was NOT sent. Say what you "
            f"saw rather than answering as though you read all of it.]")


def extract(data: bytes, name: str = "", mime: str = "") -> str:
    """Document bytes -> text, or raise with a sentence the operator can act on.

    MECHANICAL extraction only — this is the document's own text layer. When
    it raises Unextractable the document has no text layer, and `extract_best`
    is the one that then tries reading the pixels.

    The returned text is capped and SAYS SO when it was cut, so neither the
    model nor the operator can mistake a slice for the whole document.
    """
    kind = _sniff(data, name, mime)
    if kind == "pdf":
        text = _pdf(data)
    elif kind == "docx":
        text = _docx(data)
    elif kind == "zip":
        raise Unsupported(
            "this is a zip-style container Nova cannot read (only .docx is "
            "supported among those)")
    else:
        text = _plain(data)
    return _cap(text, name)


async def extract_best(data: bytes, name: str = "", mime: str = "",
                       *, allow_ocr: bool = True) -> tuple[str, str]:
    """The best text available for these bytes, WITH the source that produced it.

    Returns (text, source) where source is "mechanical" or "ocr". The source
    is returned rather than logged because the two are different kinds of
    claim — `mechanical` IS the document's text, `ocr` is a machine's reading
    of a picture of it, right in the ordinary case and wrong in specific
    recognisable ways. A caller that stores one in a field labelled "the
    document" and forgets which it was has built something that will answer
    confidently in November from a misread digit. See ocr.py.

    The chain is ordered by trust, not by convenience: the text layer wins
    whenever there is one, and OCR only runs when there is nothing to read.
    """
    is_pdf = _sniff(data, name, mime) == "pdf"
    try:
        return extract(data, name, mime), "mechanical"
    except Unextractable as mech_err:
        if not allow_ocr:
            raise
        from app import ocr
        try:
            result = await ocr.read(data, is_pdf=is_pdf)
        except ocr.OcrUnavailable as e:
            # Say BOTH halves. "No text layer" is about the document and "no
            # OCR installed" is about this server, and an operator who is
            # told only the first goes looking for a better copy of a file
            # that was fine.
            raise Unextractable(f"{mech_err} (OCR is unavailable here: {e})") from e
        except ocr.OcrEmpty as e:
            raise Unextractable(f"{mech_err} — and {e}") from e
        text = result.text
        if not result.complete:
            # the same rule as truncation: a gap the model was never shown
            # is a gap it cannot flag
            text += (f"\n\n[PARTIAL — OCR read {result.pages_read} of "
                     f"{result.pages_total} page(s) of {name or 'this file'}. "
                     f"The remaining pages were NOT read.]")
        return _cap(text, name), "ocr"


async def extract_image(data: bytes, name: str = "") -> tuple[str, str]:
    """Text out of a photograph, or raise. Same contract as extract_best.

    Separate entry point because an image has no mechanical text layer to
    try first — there is nothing to fall back FROM.
    """
    from app import ocr
    try:
        result = await ocr.read(data, is_pdf=False)
    except ocr.OcrUnavailable as e:
        raise Unsupported(f"this server cannot read text out of images: {e}") from e
    except ocr.OcrEmpty as e:
        raise Unextractable(str(e)) from e
    return _cap(result.text, name), "ocr"
